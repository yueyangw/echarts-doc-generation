import base64
import json
import os
from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from flask import Flask, request, send_from_directory

app = Flask(__name__)


@app.route("/1")
def page1():
    html = open("pie-roseType-simple.html", "r")
    return html.read()


@app.route("/2")
def page2():
    html = open("area-stack.html", "r")
    return html.read()


@app.route("/savedoc", methods=['POST'])
def save_as_document():
    datas = request.get_data()
    datas = json.loads(datas)
    img = base64_to_image(datas['imgdata'])
    template = Document('./template.docx')
    doc = doc_builder(template, img, datas['echartsData'])
    doc.save(r'report.docx')
    dic = os.getcwd()
    res = send_from_directory(dic, 'report.docx', as_attachment=True)
    return res


def base64_to_image(imgdata):
    strs = imgdata[22:]  # 返回的base64字符串前面有几个多余字符，这里删掉
    imgdata = base64.b64decode(strs)  # 把base64的图片解码成二进制返回
    return imgdata


def doc_builder(doc, img, datas):
    basic_info = datas['series'][0]
    # mapp通过type关键字识别图表类型
    mapp = {
        'pie': '饼图',
        'line': '折线图',
        'bar': '柱状图',
        'scatter': '散点图'
    }
    try:
        title = datas['title']['text']
    except:
        title = basic_info['name']
    datas = {
        'topic': title,
        'picture': img,
        'category': mapp[basic_info['type']],
        'data': basic_info['data']
    }
    # 替换标题，图片，类别（可以按需增加）
    replace_text(doc, "<topic>", datas['topic'])
    replace_text(doc, "<picture>", img)
    replace_text(doc, "<category>", datas['category'])
    return doc


'''
replace_text:
doc是要使用的docx模版
oldText为要更改的旧文本
newText为替换的新文本（当old为picture时，new为二进制图片流）
'''


def replace_text(doc, old_text, newText):
    for para in doc.paragraphs:
        if old_text in para.text:
            if old_text == '<picture>':
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                para.text = para.text.replace(old_text, 'pppiiictturree')
                for run in para.runs:
                    if 'pppiiictturree' in run.text:
                        run.text = run.text.replace('pppiiictturree', '')
                        run.add_picture(BytesIO(newText), width=Cm(14))
            else:
                para.text = para.text.replace(old_text, newText)
                for run in para.runs:
                    # 如果是标题，替换为宋体字号为27
                    if old_text == '<topic>':
                        run.font.name = '宋体'
                        run.font.size = Pt(27)
                    else:
                        run.font.name = '楷体'
                        run.font.size = Pt(16)
